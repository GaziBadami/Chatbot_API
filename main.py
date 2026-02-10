"""
main.py
Core app — /chat, /upload endpoints
Integrates with Finanvo authentication system
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from database import store_message, get_chat_history
from chatbot import get_ai_response
from database_Conv import ensure_active_conversation, update_label_if_default
import uvicorn
import os
import shutil
import fitz  # PyMuPDF
import base64
from uuid import uuid4
from docx import Document
from typing import Optional
from spire.doc import Document as SpireDocument
import json

# Import routers
from api import router as conversations_router
from admin_api import router as admin_router

app = FastAPI(title="Technowire AI Assistant")

# ============================================
# SETUP
# ============================================

UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Register routers
app.include_router(conversations_router)
app.include_router(admin_router)


# ============================================
# AUTHENTICATION HELPER
# ============================================

def extract_user_id_from_token(authorization: str) -> str:
    """
    Extract user_id from the authorization token.
    The token is the unique identifier for each user.
    """
    if not authorization:
        raise HTTPException(
            status_code=401, 
            detail="Authorization token is required. Please login first."
        )
    
    # Use the token itself as user_id (it's already unique per user)
    # Clean it up if needed (remove "Bearer " prefix if present)
    token = authorization.replace("Bearer ", "").strip()
    
    if not token:
        raise HTTPException(
            status_code=401, 
            detail="Invalid authorization token."
        )
    
    return token


# ============================================
# REQUEST MODELS
# ============================================

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[int] = None


# ============================================
# HELPER FUNCTIONS
# ============================================

def extract_text_from_pdf(file_path):
    """Extracts text content from a PDF file."""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return None

def extract_text_from_word(file_path):
    """Extracts text from .doc and .docx files."""
    try:
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        doc = SpireDocument()
        doc.LoadFromFile(file_path)
        text = doc.GetText()
        doc.Close()
        return text
    except Exception as e:
        print(f"Error extracting Word text: {e}")
        return ""

def extract_text_from_plain_file(file_path):
    """Extracts text from .txt, .json, .csv, .py files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"Text Extraction Error: {e}")
        return None

def generate_chat_label(user_message: str, ai_reply: str) -> str:
    """
    Auto-generate a short chat title from the first message
    Uses Groq to create a 4-5 word title
    Falls back to first 40 chars of user message if Groq fails
    """
    try:
        from groq import Groq
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "Generate a very short title (3-5 words max) for this conversation. Return ONLY the title, nothing else. No quotes, no punctuation."
                },
                {
                    "role": "user",
                    "content": f"User said: {user_message[:200]}\nAssistant replied: {ai_reply[:200]}"
                }
            ],
            temperature=0.3,
            max_tokens=20
        )

        label = response.choices[0].message.content.strip()
        return label[:50] if label else user_message[:40]

    except Exception as e:
        print(f"Auto-label generation failed: {e}")
        return user_message[:40]


# ============================================
# ROOT
# ============================================

@app.get("/")
async def root():
    return {"status": "online", "message": "Technowire Z-Bot API is running. Use /docs for API testing."}


# ============================================
# CHAT ENDPOINT
# ============================================

@app.post("/chat")
async def chat(
    request: ChatRequest,
    authorization: Optional[str] = Header(None)
):
    """
    Main chat endpoint - Authenticated
    
    Requires Authorization header with the user's token from Finanvo login.
    Each user gets isolated conversations based on their unique token.
    """
    try:
        # Extract user_id from Authorization header
        user_id = extract_user_id_from_token(authorization)
        
        print(f"[CHAT] User Token: {user_id[:20]}..., Message: {request.message[:50]}...")
        
        # Get or create active conversation FOR THIS USER ONLY
        active_conv = ensure_active_conversation(user_id)
        
        if not active_conv:
            raise HTTPException(
                status_code=503, 
                detail="Database connection failed. Please try again later."
            )
        
        conversation_id = active_conv['id']
        print(f"[CHAT] Conversation ID: {conversation_id}")

        # Store user message
        store_message(user_id, "user", request.message, conversation_id)

        # Get history from THIS USER's conversation only
        history = get_chat_history(user_id, conversation_id=conversation_id, limit=15)
        print(f"[CHAT] History length: {len(history)}")

        # Get AI response
        ai_reply = get_ai_response(request.message, history)

        # Store AI reply
        store_message(user_id, "assistant", ai_reply, conversation_id)

        # Auto-label
        if active_conv.get('label') == 'New Chat':
            new_label = generate_chat_label(request.message, ai_reply)
            update_label_if_default(conversation_id, new_label)
            print(f"[CHAT] Auto-labeled: {new_label}")

        return {
            "status": "success",
            "reply": ai_reply,
            "conversation_id": conversation_id,
            "user_id": user_id[:20] + "..."  # Return partial token for privacy
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[CHAT ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# ============================================
# UPLOAD ENDPOINT
# ============================================

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(None)
):
    """
    File upload endpoint - Authenticated
    
    Requires Authorization header with the user's token from Finanvo login.
    Stores file context as a message linked to user's active conversation.
    """
    try:
        # Extract user_id from Authorization header
        user_id = extract_user_id_from_token(authorization)
        
        print(f"[UPLOAD] User Token: {user_id[:20]}..., File: {file.filename}")
        
        # Ensure active conversation FOR THIS USER
        active_conv = ensure_active_conversation(user_id)
        
        if not active_conv:
            raise HTTPException(
                status_code=503, 
                detail="Database connection failed. Please try again later."
            )
        
        conversation_id = active_conv['id']

        file_extension = os.path.splitext(file.filename)[1].lower()
        unique_filename = f"{uuid4()}{file_extension}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        file_url = f"https://chatbot-api-dtad.onrender.com/uploads/{unique_filename}"

        # --- PDF ---
        if file_extension == ".pdf":
            pdf_text = extract_text_from_pdf(file_path)
            if pdf_text:
                context_msg = (
                    f"BACKGROUND_DATA: The user has uploaded a file named {file.filename}. "
                    f"Content summary: {pdf_text[:2000]}. Use this info ONLY if asked."
                )
                store_message(user_id, "user", context_msg, conversation_id)

        # --- Word Document ---
        elif file_extension in [".doc", ".docx"]:
            word_text = extract_text_from_word(file_path)
            if word_text:
                context_msg = f"SYSTEM: User uploaded a Word doc: {file.filename}. Content: {word_text[:2000]}"
                store_message(user_id, "user", context_msg, conversation_id)

        # --- Images ---
        elif file_extension in [".jpg", ".jpeg", ".png"]:
            context_msg = f"SYSTEM: User uploaded an image: {file.filename}. (The image is accessible at {file_url})"
            store_message(user_id, "user", context_msg, conversation_id)

        # --- Text / JSON / CSV / Python ---
        elif file_extension in [".json", ".txt", ".csv", ".py"]:
            extracted_content = extract_text_from_plain_file(file_path)
            if extracted_content:
                context_msg = f"BACKGROUND_DATA: Content of {file.filename}:\n{extracted_content[:2000]}"
                store_message(user_id, "user", context_msg, conversation_id)

        # Visible log message in chat
        store_message(user_id, "user", f"[File Uploaded: {file.filename}]", conversation_id)

        return {
            "status": "success",
            "file_url": file_url,
            "conversation_id": conversation_id,
            "message": f"I've received {file.filename}. Note: If I can't see the image content yet, please describe what you need me to analyze in it!"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[UPLOAD ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


# ============================================
# DEBUG ENDPOINT (Optional)
# ============================================

@app.get("/test-db")
async def test_db():
    """Test database connection"""
    from database import get_db_connection
    
    conn = get_db_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT 1 as test")
            result = cursor.fetchone()
            cursor.close()
            conn.close()
            return {
                "status": "success", 
                "message": "Database connected successfully!", 
                "test_query": result
            }
        except Exception as e:
            return {
                "status": "error", 
                "message": f"Query failed: {str(e)}"
            }
    else:
        return {
            "status": "error", 
            "message": "Database connection failed"
        }


# ============================================
# RUN
# ============================================

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)